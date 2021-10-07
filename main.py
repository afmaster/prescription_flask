from flask import Flask, render_template, request, url_for, flash, session, redirect, abort
from flask import Response, stream_with_context
from flask_wtf import FlaskForm
from wtforms import StringField, IntegerField, BooleanField
from wtforms.validators import DataRequired
import os
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from datetime import timedelta
from docx.shared import Pt
from docx2pdf import convert
import webview




app = Flask(__name__)
app.config['SECRET_KEY'] = 'MMXXII'

webview.create_window('Prescript', app, width=450, height=500, resizable=False, frameless=True, confirm_close=True)

url = "http://127.0.0.1:5000"

class Chat_form(FlaskForm):
    patient = StringField("paciente", render_kw={'autofocus': True})
    drug = StringField("medicamento")
    amount = StringField("quantidade")
    posology = IntegerField("posologia")
    observation = StringField("obs")
    observation_2 = StringField("obs_2")
    check = BooleanField("chk")#, false_values=False)
    copies = IntegerField("posologia")


@app.route("/", methods=["POST", "GET"])
def index():
    def start_func(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias):
        pic = "ineuro.jpg"
        comprador = 'comprador.jpg'
        fornecedor = 'fornecedor.jpg'
        try:
            vias = int(vias)
        except Exception as err:
            print(f"o erro das 'vias' é: {err}")
            vias = 1

        if chk == True:
            date_today = datetime.now().strftime("%d/%m/%Y")
            agora = str(date_today)
            next_month = (datetime.now() + timedelta(days=56)).strftime('%d/%m/%Y')
            two_months = (datetime.now() + timedelta(days=84)).strftime('%d/%m/%Y')

        else:
            agora = ""
            next_month = ""
            two_months = ""

        document = Document()

        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

        def prescription(dx, i, j):

            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells_row_0 = table.rows[0].cells
            paragraph_table = hdr_cells_row_0[0].paragraphs[0]
            run = paragraph_table.add_run()
            run.add_picture(pic, width=2200000, height=700000)
            k = str(i)
            hdr_cells_row_0[1].text = """
            Receituário de Controle Especial
            %sa Via - %s.
            """ % (k, j)
            paragraph_controle = hdr_cells_row_0[1].paragraphs[0]
            run = paragraph_controle.runs
            font = run[0].font
            font.size = Pt(8)

            ######################################################

            table = document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cells_row_0 = table.rows[0].cells
            hdr_cells_row_0[0].text = """
            CNPJ 13.122.535/0001-87 - CREMERS 5503
            Rua Dr. Luis Bastos do Prado, nº 1586/B, 5º andar
            Gravataí - CEP:94010-020
            Telefone (51) 3484-1745  (51) 34841756
            """
            paragraph_cnpj = hdr_cells_row_0[0].paragraphs[0]
            run = paragraph_cnpj.runs
            font = run[0].font
            font.size = Pt(8)

            table = document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cells_row_0 = table.rows[0].cells
            hdr_cells_row_0[0].text = """
            PACIENTE: %s
            Uso Interno
            1- %s ------------------------ %s
            Tomar %s comp, VO, por dia.
            %s
            %s




            %s
            """ % (paciente, medicamento, quantidade, posologia, obs, obs_2, dx)

            #################################################
            document.add_paragraph('')

            table_8 = document.add_table(rows=1, cols=2)
            table_8.style = 'Table Grid'
            hdr_cells_row_0_8 = table_8.rows[0].cells
            paragraph_table_8 = hdr_cells_row_0_8[0].paragraphs[0]
            run = paragraph_table_8.add_run()
            run.add_picture(comprador, width=1800000, height=1300000)

            hdr_cells_row_0_8_1 = table_8.rows[0].cells
            paragraph_table_8_1 = hdr_cells_row_0_8_1[1].paragraphs[0]
            run = paragraph_table_8_1.add_run()
            run.add_picture(fornecedor, width=1800000, height=1300000)

            document.add_paragraph('')

        prescription(agora, 1, "Farmácia")
        prescription(agora, 2, "Paciente")

        if vias > 1:
            prescription(next_month, 1, "Farmácia")
            prescription(next_month, 2, "Paciente")
        else:
            pass

        if vias == 3:
            prescription(two_months, 1, "Farmácia")
            prescription(two_months, 2, "Paciente")
        else:
            pass

        document.save('prescription.docx')

    def prepopulate (paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias):
        form.patient.data = paciente
        form.drug.data = medicamento
        form.amount.data = quantidade
        form.posology.data = posologia
        form.observation.data = obs
        form.observation_2.data = obs_2
        form.check.data = chk
        form.copies.data = vias

    def printing(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias):
        start_func(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
        os.startfile('prescription.docx', 'print')
        prepopulate(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)

    def opening(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias):
        start_func(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
        prepopulate(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
        try:
            os.startfile('prescription.docx')
        except Exception as err:
            print(err)


    def clear_inputs(chk, vias):
        form.patient.data = ""
        form.drug.data = ""
        form.amount.data = ""
        form.posology.data = ""
        form.observation.data = ""
        form.observation_2.data = ""
        form.check.data = chk
        form.copies.data = vias
        return render_template('index.html', form=form)

    def create_pdf(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias):
        start_func(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
        convert("prescription.docx")
        os.startfile('prescription.pdf')
        prepopulate(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)

    form = Chat_form()


    paciente = form.patient.data
    print(f"paciente = {paciente}")
    medicamento = form.drug.data
    print(f"medicamento = {medicamento}")
    quantidade= form.amount.data
    print(f"quantidade = {quantidade}")
    posologia = form.posology.data
    print(f"posoloia = {posologia}")
    obs = form.observation.data
    print(f"obs = {obs}")
    obs_2 = form.observation_2.data
    print(f"obs_2 = {obs_2}")
    chk = form.check.data
    print(f"chk = {chk}")
    vias = form.copies.data
    print(f"vias = {vias}")

    #https://stackoverflow.com/questions/43811779/use-many-submit-buttons-in-the-same-form

    print(request.form, type(request.form))
    if 'clear' in str(request.form):
        # if request.form.action == "clear":
        clear_inputs(chk, vias)
    elif 'pdf' in str(request.form):
        #elif request.form.action == "pdf":
        create_pdf(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
    elif 'print' in str(request.form):
        #elif request.form.action == "print":
        printing(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
    elif 'open' in str(request.form):
        print("identificou a afunção open")
        opening(paciente, medicamento, quantidade, posologia, obs, obs_2, chk, vias)
    else:
        return render_template('index.html', form=form)

    return render_template('index.html', form=form)



if __name__ == '__main__':
    #app.run()
    webview.start()


